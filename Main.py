import docx
import re
from docx import Document
from docx.enum.text import WD_UNDERLINE



def GetStringValue(pattern):
    # Regex re = new Regex(reStr);
    # Match match = re.Match(docText);
    #pattern='安全文明施工费(.*?)万元'
    result = re.findall(pattern, mainText)

        #result=re.findall('1<(.*)>3','1<kesagsdgy>3')
    if(result!=None):
        print(result[1])
    #print(result)
    return result


    #招标编号为:         的
    #head:      value   tail
def FillContent(head,tail,value):#查找匹配输入

    # head='招标编号为'
    # tail='的'
    # value='afsdgasdjkghasjh'

    start = False
    for p in doc.paragraphs:
        if ('（封面）' in p.text):
            start = True
        if('投标工程名称：' in p.text):
            print(1)
        if (start):
            # if ('投标文件内容' in p.text):
            #     start = False
            # print(p.text)
            for i in range(len(p.runs)):
                r = p.runs[i]
                lastText = None#todoooooooooooooooooooooooooooooooooooooooooo
                nextText = None
                pLenth=len(p.runs)
                if (i != 0):
                    plusMax = 2
                    plus = 1
                    lastText = ''
                    while (i - plus >= 0 and plus <= plusMax):
                        if(p.runs[i-plus].text.isspace()):
                            plus=plus+1
                            plusMax=plusMax+1
                            continue
                        else:
                            lastText = lastText + p.runs[i - plus].text
                            plus = plus + 1
                    #lastText = p.runs[i - 1].text+p.runs[i].text
                if (i != pLenth - 1):#循环取后面的字符串，遇到空格跳过,
                    plusMax=2
                    plus=1
                    nextText=''
                    while(i+plus<pLenth and plus<=plusMax):
                        if (p.runs[i + plus].text.isspace()):
                            plus=plus+1
                            plusMax=plusMax+1#只匹配不是空格的字符串
                            continue
                        else:
                            nextText = nextText +  p.runs[i + plus].text
                            plus=plus+1

                #if(lastText==None or nextText==None):
                    #print('nnnnnnnnnnnnnnnnnnnnnnnnnnn')
                headMatch=False
                tailMatch=False
                if(head==''):
                    headMatch=True
                else:
                    if(lastText):
                        if(head in lastText):
                           headMatch=True

                if(tail==''):
                    tailMatch=True
                else:
                    if(nextText):
                        if(tail in nextText):
                            tailMatch=True#end todooooooooooooooooooooooooooooooooooooo

                if ( headMatch and tailMatch):
                    # print(r.underline)
                    # print(r.text)

                    # 输入数据
                    if ((r.underline == WD_UNDERLINE.THICK or r.underline == True) and len(r.text) > 3):
                        input = True
                        for char in r.text:
                            if char != ' ':
                                input = False
                        if (input):
                            r.clear()

                            r.add_text(value)
                            # r.text='afsdfa'
                            print(value)

if __name__ == "__main__":

    path = r"C:\Users\Administrator\Desktop\Test\test.docx"
    doc = Document(path)
    mainText = ''
    for p in doc.paragraphs:
        mainText = mainText + p.text
    for t in doc.tables:
        for cell in t._cells:
            mainText = mainText + cell.text

    #print(mainText)
    安全文明施工费 = GetStringValue( '安全文明施工费(.*?)万元')
    #工期 = GetStringValue( @ "工期要求\a(?<key>.*?)日历天")
    # 担保金百分比 = GetStringValue( @ "交纳中标价(?<key>.*?)%的履约保证金")
    # 担保金值 = GetStringValue( @ "投标担保金额\a(?<key>.*?)元")
    # 日期 = GetStringValue( @ "投标文件递交截止时间.*?2019年(?<key>.*?)日")
    # 误期违约金额 = GetStringValue( @ "误期违约金额.*?(?<key>.*?)元")
    # 预付款金额 = GetStringValue( @ "预付款金额(:|：).*?(?<key>.*?)预付款保函金额")

    #FillContent('a','b','c')
    FillContent('到贵方的招标编号为','的','gasjgs')
    FillContent('的','工程','52345')
    FillContent('投标工程名称','','fffffff')
    doc.save(r"C:\Users\Administrator\Desktop\Test\testResult.docx")
